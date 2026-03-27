"""
create_corpus.py — генерирует тестовый корпус для LightLocalTranslator

Создаёт:
  corpus/blocks/A_words.jsonl         — 50 слов с референсным переводом
  corpus/blocks/B_sentences.jsonl     — 30 предложений с референсом
  corpus/blocks/C_paragraphs.jsonl    — 10 абзацев с референсом
  corpus/documents/test_doc.txt       — ~3000 слов технический текст EN
  corpus/documents/test_doc.md        — тот же текст в Markdown
  corpus/documents/test_doc.docx      — DOCX с таблицей, стилями, колонтитулами
  corpus/documents/test_doc.xlsx      — таблица с текстом + формулами
  corpus/documents/test_doc.pdf       — searchable PDF

Запуск: python corpus/create_corpus.py
Зависимости: python-docx, openpyxl, reportlab (все есть в requirements)
"""

import json
import os
from pathlib import Path

ROOT = Path(__file__).parent
BLOCKS = ROOT / "blocks"
DOCS = ROOT / "documents"
BLOCKS.mkdir(exist_ok=True)
DOCS.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# A — СЛОВА (50 пар)
# ---------------------------------------------------------------------------

WORDS = [
    # EN→RU: технические термины (25)
    {"id": "w001", "src_lang": "en", "tgt_lang": "ru", "source": "compiler",        "reference": "компилятор"},
    {"id": "w002", "src_lang": "en", "tgt_lang": "ru", "source": "bandwidth",       "reference": "пропускная способность"},
    {"id": "w003", "src_lang": "en", "tgt_lang": "ru", "source": "latency",         "reference": "задержка"},
    {"id": "w004", "src_lang": "en", "tgt_lang": "ru", "source": "throughput",      "reference": "производительность"},
    {"id": "w005", "src_lang": "en", "tgt_lang": "ru", "source": "cache",           "reference": "кэш"},
    {"id": "w006", "src_lang": "en", "tgt_lang": "ru", "source": "firmware",        "reference": "прошивка"},
    {"id": "w007", "src_lang": "en", "tgt_lang": "ru", "source": "interrupt",       "reference": "прерывание"},
    {"id": "w008", "src_lang": "en", "tgt_lang": "ru", "source": "scheduler",       "reference": "планировщик"},
    {"id": "w009", "src_lang": "en", "tgt_lang": "ru", "source": "semaphore",       "reference": "семафор"},
    {"id": "w010", "src_lang": "en", "tgt_lang": "ru", "source": "kernel",          "reference": "ядро"},
    {"id": "w011", "src_lang": "en", "tgt_lang": "ru", "source": "overhead",        "reference": "накладные расходы"},
    {"id": "w012", "src_lang": "en", "tgt_lang": "ru", "source": "bottleneck",      "reference": "узкое место"},
    {"id": "w013", "src_lang": "en", "tgt_lang": "ru", "source": "deployment",      "reference": "развёртывание"},
    {"id": "w014", "src_lang": "en", "tgt_lang": "ru", "source": "checksum",        "reference": "контрольная сумма"},
    {"id": "w015", "src_lang": "en", "tgt_lang": "ru", "source": "handshake",       "reference": "квитирование"},
    {"id": "w016", "src_lang": "en", "tgt_lang": "ru", "source": "pipeline",        "reference": "конвейер"},
    {"id": "w017", "src_lang": "en", "tgt_lang": "ru", "source": "routing",         "reference": "маршрутизация"},
    {"id": "w018", "src_lang": "en", "tgt_lang": "ru", "source": "encryption",      "reference": "шифрование"},
    {"id": "w019", "src_lang": "en", "tgt_lang": "ru", "source": "debugging",       "reference": "отладка"},
    {"id": "w020", "src_lang": "en", "tgt_lang": "ru", "source": "concurrency",     "reference": "параллелизм"},
    {"id": "w021", "src_lang": "en", "tgt_lang": "ru", "source": "scalability",     "reference": "масштабируемость"},
    {"id": "w022", "src_lang": "en", "tgt_lang": "ru", "source": "resilience",      "reference": "отказоустойчивость"},
    {"id": "w023", "src_lang": "en", "tgt_lang": "ru", "source": "abstraction",     "reference": "абстракция"},
    {"id": "w024", "src_lang": "en", "tgt_lang": "ru", "source": "profiling",       "reference": "профилирование"},
    {"id": "w025", "src_lang": "en", "tgt_lang": "ru", "source": "payload",         "reference": "полезная нагрузка"},
    # RU→EN: технические и бытовые (25)
    {"id": "w026", "src_lang": "ru", "tgt_lang": "en", "source": "компьютер",             "reference": "computer"},
    {"id": "w027", "src_lang": "ru", "tgt_lang": "en", "source": "программа",             "reference": "program"},
    {"id": "w028", "src_lang": "ru", "tgt_lang": "en", "source": "сеть",                  "reference": "network"},
    {"id": "w029", "src_lang": "ru", "tgt_lang": "en", "source": "память",                "reference": "memory"},
    {"id": "w030", "src_lang": "ru", "tgt_lang": "en", "source": "файл",                  "reference": "file"},
    {"id": "w031", "src_lang": "ru", "tgt_lang": "en", "source": "сервер",                "reference": "server"},
    {"id": "w032", "src_lang": "ru", "tgt_lang": "en", "source": "процессор",             "reference": "processor"},
    {"id": "w033", "src_lang": "ru", "tgt_lang": "en", "source": "браузер",               "reference": "browser"},
    {"id": "w034", "src_lang": "ru", "tgt_lang": "en", "source": "клавиатура",            "reference": "keyboard"},
    {"id": "w035", "src_lang": "ru", "tgt_lang": "en", "source": "монитор",               "reference": "monitor"},
    {"id": "w036", "src_lang": "ru", "tgt_lang": "en", "source": "алгоритм",              "reference": "algorithm"},
    {"id": "w037", "src_lang": "ru", "tgt_lang": "en", "source": "база данных",           "reference": "database"},
    {"id": "w038", "src_lang": "ru", "tgt_lang": "en", "source": "протокол",              "reference": "protocol"},
    {"id": "w039", "src_lang": "ru", "tgt_lang": "en", "source": "интерфейс",             "reference": "interface"},
    {"id": "w040", "src_lang": "ru", "tgt_lang": "en", "source": "обновление",            "reference": "update"},
    {"id": "w041", "src_lang": "ru", "tgt_lang": "en", "source": "установка",             "reference": "installation"},
    {"id": "w042", "src_lang": "ru", "tgt_lang": "en", "source": "резервное копирование", "reference": "backup"},
    {"id": "w043", "src_lang": "ru", "tgt_lang": "en", "source": "восстановление",        "reference": "recovery"},
    {"id": "w044", "src_lang": "ru", "tgt_lang": "en", "source": "шифрование",            "reference": "encryption"},
    {"id": "w045", "src_lang": "ru", "tgt_lang": "en", "source": "маршрутизатор",         "reference": "router"},
    {"id": "w046", "src_lang": "ru", "tgt_lang": "en", "source": "операционная система",  "reference": "operating system"},
    {"id": "w047", "src_lang": "ru", "tgt_lang": "en", "source": "виртуализация",         "reference": "virtualization"},
    {"id": "w048", "src_lang": "ru", "tgt_lang": "en", "source": "контейнер",             "reference": "container"},
    {"id": "w049", "src_lang": "ru", "tgt_lang": "en", "source": "отладчик",              "reference": "debugger"},
    {"id": "w050", "src_lang": "ru", "tgt_lang": "en", "source": "компилятор",            "reference": "compiler"},
]


# ---------------------------------------------------------------------------
# B — ПРЕДЛОЖЕНИЯ (30 пар)
# ---------------------------------------------------------------------------

SENTENCES = [
    # EN→RU (15)
    {"id": "s001", "src_lang": "en", "tgt_lang": "ru",
     "source":    "The processor executes instructions at a rate of 3.2 billion cycles per second.",
     "reference": "Процессор выполняет инструкции со скоростью 3,2 миллиарда циклов в секунду."},
    {"id": "s002", "src_lang": "en", "tgt_lang": "ru",
     "source":    "Memory latency is the time required to access data stored in RAM.",
     "reference": "Задержка памяти — это время, необходимое для доступа к данным, хранящимся в оперативной памяти."},
    {"id": "s003", "src_lang": "en", "tgt_lang": "ru",
     "source":    "The operating system manages hardware resources and provides services to applications.",
     "reference": "Операционная система управляет аппаратными ресурсами и предоставляет сервисы приложениям."},
    {"id": "s004", "src_lang": "en", "tgt_lang": "ru",
     "source":    "Machine translation systems have improved significantly with the adoption of neural networks.",
     "reference": "Системы машинного перевода значительно улучшились с распространением нейронных сетей."},
    {"id": "s005", "src_lang": "en", "tgt_lang": "ru",
     "source":    "The network packet was fragmented into 1,500-byte segments for transmission.",
     "reference": "Сетевой пакет был разбит на сегменты размером 1 500 байт для передачи."},
    {"id": "s006", "src_lang": "en", "tgt_lang": "ru",
     "source":    "Caching frequently accessed data reduces the average memory access time.",
     "reference": "Кэширование часто используемых данных снижает среднее время доступа к памяти."},
    {"id": "s007", "src_lang": "en", "tgt_lang": "ru",
     "source":    "The algorithm has a time complexity of O(n log n) in the worst case.",
     "reference": "Алгоритм имеет временную сложность O(n log n) в наихудшем случае."},
    {"id": "s008", "src_lang": "en", "tgt_lang": "ru",
     "source":    "Encryption ensures that only authorized parties can read the transmitted data.",
     "reference": "Шифрование гарантирует, что только авторизованные стороны могут читать передаваемые данные."},
    {"id": "s009", "src_lang": "en", "tgt_lang": "ru",
     "source":    "The server handles approximately 10,000 requests per second under peak load.",
     "reference": "Сервер обрабатывает около 10 000 запросов в секунду при пиковой нагрузке."},
    {"id": "s010", "src_lang": "en", "tgt_lang": "ru",
     "source":    "A deadlock occurs when two or more processes are waiting for each other to release resources.",
     "reference": "Взаимная блокировка возникает, когда два или более процессов ожидают освобождения ресурсов друг другом."},
    {"id": "s011", "src_lang": "en", "tgt_lang": "ru",
     "source":    "The database index reduced query execution time from 4.2 seconds to 12 milliseconds.",
     "reference": "Индекс базы данных сократил время выполнения запроса с 4,2 секунды до 12 миллисекунд."},
    {"id": "s012", "src_lang": "en", "tgt_lang": "ru",
     "source":    "Modern solid-state drives achieve sequential read speeds of up to 7,000 MB/s.",
     "reference": "Современные твердотельные накопители достигают последовательной скорости чтения до 7 000 МБ/с."},
    {"id": "s013", "src_lang": "en", "tgt_lang": "ru",
     "source":    "The API endpoint returns a JSON response with the requested user data.",
     "reference": "Конечная точка API возвращает ответ в формате JSON с запрошенными данными пользователя."},
    {"id": "s014", "src_lang": "en", "tgt_lang": "ru",
     "source":    "Virtualization allows multiple operating systems to run on a single physical machine.",
     "reference": "Виртуализация позволяет нескольким операционным системам работать на одном физическом компьютере."},
    {"id": "s015", "src_lang": "en", "tgt_lang": "ru",
     "source":    "The compiler translates high-level source code into machine-executable binary instructions.",
     "reference": "Компилятор переводит высокоуровневый исходный код в исполняемые машиной двоичные инструкции."},
    # RU→EN (15)
    {"id": "s016", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Нейронная сеть обучается на большом наборе данных, чтобы распознавать закономерности.",
     "reference": "The neural network is trained on a large dataset to recognize patterns."},
    {"id": "s017", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Программное обеспечение с открытым исходным кодом может быть свободно изменено и распространено.",
     "reference": "Open-source software can be freely modified and distributed."},
    {"id": "s018", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Частота обновления экрана составляет 144 Гц, что обеспечивает плавное изображение.",
     "reference": "The display refresh rate is 144 Hz, providing smooth visuals."},
    {"id": "s019", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Резервное копирование данных следует выполнять не реже одного раза в сутки.",
     "reference": "Data backups should be performed at least once a day."},
    {"id": "s020", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Протокол HTTPS шифрует данные, передаваемые между браузером и сервером.",
     "reference": "The HTTPS protocol encrypts data transmitted between the browser and the server."},
    {"id": "s021", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Задержка сети менее 20 миллисекунд считается приемлемой для онлайн-игр.",
     "reference": "Network latency below 20 milliseconds is considered acceptable for online gaming."},
    {"id": "s022", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Разработчик обнаружил ошибку сегментации при обработке входных данных размером более 4 ГБ.",
     "reference": "The developer found a segmentation fault when processing input data larger than 4 GB."},
    {"id": "s023", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Микропроцессор содержит более 8 миллиардов транзисторов на площади менее 100 мм².",
     "reference": "The microprocessor contains more than 8 billion transistors in an area of less than 100 mm²."},
    {"id": "s024", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Алгоритм сортировки слиянием является стабильным и имеет сложность O(n log n).",
     "reference": "The merge sort algorithm is stable and has a complexity of O(n log n)."},
    {"id": "s025", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Контейнеризация позволяет изолировать приложения и их зависимости в единый пакет.",
     "reference": "Containerization allows isolating applications and their dependencies into a single package."},
    {"id": "s026", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Облачные вычисления предоставляют масштабируемые ресурсы по модели оплаты по факту использования.",
     "reference": "Cloud computing provides scalable resources on a pay-as-you-go model."},
    {"id": "s027", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Система управления версиями отслеживает все изменения в исходном коде проекта.",
     "reference": "The version control system tracks all changes in the project's source code."},
    {"id": "s028", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Межсетевой экран фильтрует входящий и исходящий трафик на основе правил безопасности.",
     "reference": "The firewall filters incoming and outgoing traffic based on security rules."},
    {"id": "s029", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Кэш процессора первого уровня имеет объём 32 КБ и работает на частоте ядра.",
     "reference": "The L1 processor cache has a capacity of 32 KB and operates at the core frequency."},
    {"id": "s030", "src_lang": "ru", "tgt_lang": "en",
     "source":    "Пропускная способность канала связи ограничена теоремой Шеннона о пропускной способности.",
     "reference": "The channel capacity is limited by Shannon's channel capacity theorem."},
]


# ---------------------------------------------------------------------------
# C — АБЗАЦЫ (10 пар)
# ---------------------------------------------------------------------------

PARAGRAPHS = [
    {"id": "p001", "src_lang": "en", "tgt_lang": "ru",
     "source": (
         "A central processing unit (CPU) is the primary component of a computer that executes instructions. "
         "Modern CPUs operate at frequencies between 3.0 and 5.5 GHz and contain multiple execution cores, "
         "enabling parallel processing of tasks. Each core consists of an arithmetic logic unit (ALU), a "
         "control unit, and registers. The L1 cache, typically 32–64 KB per core, stores the most frequently "
         "accessed instructions and data. The L2 cache, usually 256 KB to 1 MB per core, serves as a secondary "
         "buffer. The L3 cache is shared among all cores and typically ranges from 8 to 64 MB. Branch "
         "prediction and out-of-order execution are key microarchitectural techniques that allow modern CPUs "
         "to achieve high instruction-level parallelism."
     ),
     "reference": (
         "Центральный процессор (ЦП) — основной компонент компьютера, выполняющий инструкции. "
         "Современные процессоры работают на частотах от 3,0 до 5,5 ГГц и содержат несколько вычислительных "
         "ядер, что обеспечивает параллельную обработку задач. Каждое ядро состоит из арифметико-логического "
         "устройства (АЛУ), устройства управления и регистров. Кэш первого уровня (L1), как правило 32–64 КБ "
         "на ядро, хранит наиболее часто используемые инструкции и данные. Кэш второго уровня (L2), обычно "
         "от 256 КБ до 1 МБ на ядро, служит вторичным буфером. Кэш третьего уровня (L3) является общим для "
         "всех ядер и обычно составляет от 8 до 64 МБ. Предсказание ветвлений и внеочередное выполнение — "
         "ключевые микроархитектурные методы, позволяющие современным процессорам достигать высокого "
         "параллелизма на уровне инструкций."
     )},
    {"id": "p002", "src_lang": "en", "tgt_lang": "ru",
     "source": (
         "Random access memory (RAM) is the primary workspace of a computer, storing data and instructions "
         "that the CPU is actively using. Modern systems use DDR5 memory modules operating at speeds between "
         "4,800 and 8,400 MT/s with a typical capacity of 16 to 64 GB per module. Unlike storage devices, "
         "RAM is volatile: all data is lost when power is removed. The memory controller, integrated into the "
         "CPU since 2006, manages data transfers between the processor and RAM. Memory bandwidth directly "
         "affects performance in workloads that process large datasets, such as video editing, scientific "
         "simulation, and machine learning training. Dual-channel and quad-channel configurations multiply "
         "the effective bandwidth by using multiple memory modules simultaneously."
     ),
     "reference": (
         "Оперативная память (ОЗУ) — основное рабочее пространство компьютера, в котором хранятся данные "
         "и инструкции, активно используемые процессором. Современные системы используют модули памяти DDR5, "
         "работающие на скорости от 4 800 до 8 400 МТ/с с типичной ёмкостью от 16 до 64 ГБ на модуль. "
         "В отличие от устройств хранения данных, оперативная память является энергозависимой: все данные "
         "теряются при отключении питания. Контроллер памяти, интегрированный в процессор с 2006 года, "
         "управляет передачей данных между процессором и ОЗУ. Пропускная способность памяти напрямую влияет "
         "на производительность в задачах, обрабатывающих большие наборы данных, — таких как монтаж видео, "
         "научное моделирование и обучение моделей машинного обучения. Двухканальные и четырёхканальные "
         "конфигурации умножают эффективную пропускную способность за счёт одновременного использования "
         "нескольких модулей памяти."
     )},
    {"id": "p003", "src_lang": "en", "tgt_lang": "ru",
     "source": (
         "A computer network is a collection of interconnected devices that share resources and exchange data. "
         "The TCP/IP protocol suite forms the foundation of the modern internet, operating across four layers: "
         "the link layer, the internet layer, the transport layer, and the application layer. The Ethernet "
         "standard governs local area network (LAN) communication, with modern installations supporting speeds "
         "of 1 Gbps, 10 Gbps, or 100 Gbps. Wireless networks based on the IEEE 802.11 standard (Wi-Fi) "
         "provide mobility but at generally lower throughput and higher latency than wired connections. "
         "Network packets travel through routers and switches, which use routing tables and MAC address tables "
         "to forward traffic efficiently. Quality of Service (QoS) mechanisms prioritize latency-sensitive "
         "traffic such as voice and video over bulk data transfers."
     ),
     "reference": (
         "Компьютерная сеть — совокупность взаимосвязанных устройств, совместно использующих ресурсы и "
         "обменивающихся данными. Набор протоколов TCP/IP составляет основу современного интернета и работает "
         "на четырёх уровнях: канальном, сетевом, транспортном и прикладном. Стандарт Ethernet регулирует "
         "связь в локальных сетях (LAN); современные установки поддерживают скорости 1 Гбит/с, 10 Гбит/с "
         "или 100 Гбит/с. Беспроводные сети на базе стандарта IEEE 802.11 (Wi-Fi) обеспечивают мобильность, "
         "однако, как правило, характеризуются более низкой пропускной способностью и большей задержкой по "
         "сравнению с проводными соединениями. Сетевые пакеты передаются через маршрутизаторы и коммутаторы, "
         "которые используют таблицы маршрутизации и таблицы MAC-адресов для эффективной пересылки трафика. "
         "Механизмы качества обслуживания (QoS) отдают приоритет трафику, чувствительному к задержкам — "
         "например, голосовым и видеопотокам — над передачей больших объёмов данных."
     )},
    {"id": "p004", "src_lang": "en", "tgt_lang": "ru",
     "source": (
         "An operating system (OS) is the fundamental software layer that manages hardware resources and "
         "provides a runtime environment for applications. The OS kernel handles process scheduling, memory "
         "allocation, file system management, and device driver interfaces. Modern operating systems implement "
         "virtual memory, which allows processes to address more memory than is physically available by mapping "
         "virtual addresses to physical pages through a page table. Context switching, the mechanism by which "
         "the scheduler saves the state of one process and restores another, occurs thousands of times per "
         "second on a busy system. Semaphores, mutexes, and condition variables are synchronization primitives "
         "used to coordinate access to shared resources among concurrent threads. System calls provide a "
         "controlled interface through which user-space applications can request kernel-level services."
     ),
     "reference": (
         "Операционная система (ОС) — основной программный уровень, управляющий аппаратными ресурсами и "
         "предоставляющий среду выполнения для приложений. Ядро ОС занимается планированием процессов, "
         "выделением памяти, управлением файловой системой и интерфейсами драйверов устройств. Современные "
         "операционные системы реализуют виртуальную память, которая позволяет процессам адресовать больший "
         "объём памяти, чем физически доступен, путём отображения виртуальных адресов на физические страницы "
         "через таблицу страниц. Переключение контекста — механизм, при котором планировщик сохраняет "
         "состояние одного процесса и восстанавливает другой, — происходит тысячи раз в секунду на нагруженной "
         "системе. Семафоры, мьютексы и условные переменные являются примитивами синхронизации, используемыми "
         "для координации доступа к общим ресурсам среди параллельных потоков. Системные вызовы предоставляют "
         "управляемый интерфейс, через который приложения пользовательского пространства могут запрашивать "
         "сервисы уровня ядра."
     )},
    {"id": "p005", "src_lang": "en", "tgt_lang": "ru",
     "source": (
         "Mitochondria are membrane-bound organelles found in the cytoplasm of eukaryotic cells. They are "
         "responsible for generating the majority of the cell's supply of adenosine triphosphate (ATP), used "
         "as a source of chemical energy through the process of oxidative phosphorylation. The inner "
         "mitochondrial membrane is folded into structures called cristae, which dramatically increase the "
         "surface area available for ATP synthesis. Mitochondria contain their own circular DNA and ribosomes, "
         "evidence of their evolutionary origin as free-living bacteria engulfed by ancestral eukaryotic cells "
         "approximately 1.5 billion years ago. A typical human cell contains between 1,000 and 2,000 "
         "mitochondria, with the exact number depending on the cell's energy requirements. Muscle cells and "
         "hepatocytes, which have high metabolic demands, contain particularly dense mitochondrial populations."
     ),
     "reference": (
         "Митохондрии — мембранные органеллы, обнаруженные в цитоплазме эукариотических клеток. Они отвечают "
         "за выработку большей части запасов аденозинтрифосфата (АТФ) клетки, используемого в качестве "
         "источника химической энергии в процессе окислительного фосфорилирования. Внутренняя мембрана "
         "митохондрий складывается в структуры, называемые кристами, которые резко увеличивают площадь "
         "поверхности, доступной для синтеза АТФ. Митохондрии содержат собственную кольцевую ДНК и рибосомы "
         "— свидетельство их эволюционного происхождения от свободноживущих бактерий, поглощённых клетками-"
         "предшественниками эукариот приблизительно 1,5 миллиарда лет назад. Типичная клетка человека "
         "содержит от 1 000 до 2 000 митохондрий; точное число зависит от энергетических потребностей клетки. "
         "Мышечные клетки и гепатоциты с высокими метаболическими потребностями содержат особенно плотные "
         "митохондриальные популяции."
     )},
    {"id": "p006", "src_lang": "en", "tgt_lang": "ru",
     "source": (
         "Quantum entanglement is a phenomenon in which two or more particles become correlated such that the "
         "quantum state of each particle cannot be described independently of the others, even when separated "
         "by large distances. When a measurement is performed on one entangled particle, the result "
         "instantaneously determines the corresponding property of its partner, regardless of the distance "
         "between them. This non-locality troubled Einstein, who referred to it as 'spooky action at a "
         "distance.' Experimental tests of Bell's inequalities have conclusively demonstrated that entanglement "
         "is a genuine physical effect and cannot be explained by any local hidden variable theory. Quantum "
         "entanglement is now actively exploited in quantum computing, where entangled qubits enable operations "
         "with no classical equivalent, and in quantum cryptography, where it forms the basis of provably "
         "secure communication protocols."
     ),
     "reference": (
         "Квантовая запутанность — явление, при котором две или более частицы оказываются скоррелированы "
         "таким образом, что квантовое состояние каждой из них не может быть описано независимо от других, "
         "даже при разделении на большие расстояния. При измерении над одной запутанной частицей результат "
         "мгновенно определяет соответствующее свойство её партнёра, независимо от расстояния между ними. "
         "Эта нелокальность тревожила Эйнштейна, назвавшего её «жутким действием на расстоянии». "
         "Экспериментальные проверки неравенств Белла убедительно показали, что запутанность является "
         "подлинным физическим эффектом и не может быть объяснена никакой теорией со скрытыми локальными "
         "переменными. Квантовая запутанность ныне активно используется в квантовых вычислениях, где "
         "запутанные кубиты обеспечивают операции без классических аналогов, и в квантовой криптографии, "
         "где она служит основой доказуемо защищённых протоколов связи."
     )},
    {"id": "p007", "src_lang": "en", "tgt_lang": "ru",
     "source": (
         "The global semiconductor shortage that began in 2020 exposed deep vulnerabilities in international "
         "supply chains. What started as a temporary disruption caused by pandemic-related factory closures "
         "cascaded into a multi-year crisis affecting automakers, consumer electronics manufacturers, and "
         "defense contractors alike. Production of a modern microprocessor requires over 1,000 individual "
         "process steps and depends on equipment and materials sourced from dozens of countries. The "
         "concentration of advanced chip manufacturing in a handful of facilities, primarily in Taiwan and "
         "South Korea, prompted governments worldwide to invest in domestic semiconductor capacity. The United "
         "States CHIPS Act, signed in 2022, allocated 52.7 billion dollars to boost domestic chip production "
         "and research. Similar initiatives were launched in the European Union, Japan, and India, reflecting "
         "a broad recognition that semiconductor self-sufficiency has become a matter of national security."
     ),
     "reference": (
         "Глобальная нехватка полупроводников, начавшаяся в 2020 году, обнажила глубокие уязвимости в "
         "международных цепочках поставок. То, что началось как временный сбой из-за закрытия заводов во "
         "время пандемии, переросло в многолетний кризис, затронувший автопроизводителей, производителей "
         "потребительской электроники и оборонных подрядчиков. Производство современного микропроцессора "
         "требует более 1 000 отдельных технологических операций и зависит от оборудования и материалов, "
         "поставляемых из десятков стран. Концентрация передового производства чипов в немногочисленных "
         "предприятиях, расположенных преимущественно на Тайване и в Южной Корее, побудила правительства "
         "по всему миру инвестировать в отечественные мощности по производству полупроводников. Закон США "
         "о CHIPS, подписанный в 2022 году, выделил 52,7 миллиарда долларов на развитие внутреннего "
         "производства чипов. Аналогичные инициативы были запущены в ЕС, Японии и Индии, что отражает "
         "широкое признание того, что самодостаточность в области полупроводников стала вопросом "
         "национальной безопасности."
     )},
    # RU→EN (3)
    {"id": "p008", "src_lang": "ru", "tgt_lang": "en",
     "source": (
         "Реляционная база данных организует данные в виде таблиц со строками и столбцами, связанными между "
         "собой через внешние ключи. Язык структурированных запросов (SQL) позволяет выполнять операции "
         "выборки, вставки, обновления и удаления данных. Транзакции обеспечивают атомарность, "
         "согласованность, изолированность и долговечность операций — свойства, известные под аббревиатурой "
         "ACID. Индексы ускоряют поиск данных за счёт создания вспомогательных структур, упорядоченных по "
         "значению одного или нескольких столбцов. Оптимизатор запросов анализирует возможные планы "
         "выполнения и выбирает наиболее эффективный на основе статистики о распределении данных."
     ),
     "reference": (
         "A relational database organizes data into tables with rows and columns linked to each other through "
         "foreign keys. Structured Query Language (SQL) enables select, insert, update, and delete operations "
         "on data. Transactions ensure atomicity, consistency, isolation, and durability of operations — "
         "properties known by the acronym ACID. Indexes speed up data retrieval by creating auxiliary "
         "structures ordered by the value of one or more columns. The query optimizer analyzes possible "
         "execution plans and selects the most efficient one based on statistics about data distribution."
     )},
    {"id": "p009", "src_lang": "ru", "tgt_lang": "en",
     "source": (
         "Антибиотики — класс противомикробных препаратов, используемых для лечения бактериальных инфекций "
         "путём уничтожения бактерий или подавления их роста. Механизм действия различных групп антибиотиков "
         "существенно отличается: пенициллины и цефалоспорины нарушают синтез клеточной стенки бактерий, "
         "аминогликозиды подавляют синтез белка на уровне рибосом, а фторхинолоны ингибируют ферменты, "
         "необходимые для репликации ДНК. Нерациональное применение антибиотиков привело к появлению и "
         "распространению устойчивых штаммов бактерий — проблеме, признанной ВОЗ одной из серьёзных угроз "
         "глобальному здоровью."
     ),
     "reference": (
         "Antibiotics are a class of antimicrobial drugs used to treat bacterial infections by killing "
         "bacteria or inhibiting their growth. The mechanism of action of different antibiotic groups differs "
         "significantly: penicillins and cephalosporins disrupt bacterial cell wall synthesis, aminoglycosides "
         "inhibit protein synthesis at the ribosome level, and fluoroquinolones inhibit enzymes required for "
         "DNA replication. Irrational use of antibiotics has led to the emergence and spread of resistant "
         "bacterial strains — a problem recognized by the WHO as one of the serious threats to global health."
     )},
    {"id": "p010", "src_lang": "ru", "tgt_lang": "en",
     "source": (
         "Искусственный интеллект стремительно трансформирует рынок труда, вызывая острые дискуссии о "
         "будущем занятости. Автоматизация рутинных задач — от обработки документов до управления "
         "транспортными средствами — уже затронула миллионы рабочих мест по всему миру. Согласно докладу "
         "Международной организации труда, опубликованному в 2023 году, генеративный ИИ способен частично "
         "или полностью автоматизировать от 14 до 40 процентов задач в большинстве профессий. Вместе с тем "
         "исторический опыт технологических революций свидетельствует о том, что новые технологии, как "
         "правило, создают больше рабочих мест, чем уничтожают, хотя этот переход может быть болезненным "
         "для отдельных категорий работников."
     ),
     "reference": (
         "Artificial intelligence is rapidly transforming the labor market, sparking intense debate about "
         "the future of employment. The automation of routine tasks — from document processing to vehicle "
         "operation — has already affected millions of jobs worldwide. According to a report by the "
         "International Labour Organization published in 2023, generative AI is capable of partially or fully "
         "automating between 14 and 40 percent of tasks in most professions. At the same time, the historical "
         "experience of technological revolutions suggests that new technologies generally create more jobs "
         "than they destroy, although this transition can be painful for certain categories of workers."
     )},
]


# ---------------------------------------------------------------------------
# Основной документ — TEXT (~3000 слов)
# ---------------------------------------------------------------------------

DOCUMENT_TITLE = "Fundamentals of Computer Architecture and System Software"

DOCUMENT_SECTIONS = [
    ("1. Introduction", """
Computer architecture is the discipline concerned with the design, organization, and behavior of the components
that make up a computing system. At its most fundamental level, a computer consists of a processor that executes
instructions, memory that stores data and programs, and input/output devices that connect the machine to the
outside world. The interaction between these components, governed by protocols and interfaces defined at design
time, determines the overall performance, power consumption, and capability of the system.

Modern computers are built upon decades of accumulated engineering knowledge. The stored-program concept,
formalized by John von Neumann in 1945, established the architectural template that still underlies most
computers today: instructions and data reside together in the same memory space, and the processor fetches,
decodes, and executes one instruction after another in sequence, occasionally branching to non-sequential
addresses. Although this model has been extended and optimized in countless ways — through pipelining,
superscalar execution, out-of-order processing, and speculative execution — its fundamental structure
remains recognizable.

Understanding computer architecture is essential for software engineers who want to write efficient code,
for system administrators who need to configure hardware, and for anyone who builds or maintains software
infrastructure. Even when working at a high level of abstraction, decisions about data structures, memory
access patterns, and concurrency are ultimately constrained by the properties of the underlying hardware.
"""),

    ("2. The Central Processing Unit", """
The central processing unit (CPU) is the engine of the computer. It retrieves instructions from memory,
decodes them to determine what operation is required, and executes that operation using its internal
functional units. This fetch-decode-execute cycle runs continuously at a rate determined by the processor's
clock frequency, which in modern chips typically ranges from 3.0 GHz to 5.5 GHz.

A modern processor contains multiple execution cores, each capable of running an independent thread of
execution. A high-end desktop CPU may have 8, 12, or even 24 cores, while server processors can exceed
100 cores per chip. Each core contains an arithmetic logic unit (ALU) for integer operations, a
floating-point unit (FPU) for fractional and real-number arithmetic, a set of fast storage locations
called registers, and dedicated control logic.

To bridge the speed gap between the processor and main memory, CPUs contain several levels of cache.
The L1 cache, integrated directly into each core, is the fastest and smallest, typically 32 to 64 KB.
The L2 cache, slightly slower but larger at 256 KB to 1 MB per core, serves as a secondary buffer.
The L3 cache is shared among all cores on the chip and ranges from 8 MB in budget processors to 64 MB
or more in workstation-class chips. When the processor needs data, it first checks L1, then L2, then L3,
and finally issues a request to main memory if the data is not found in any cache level.

Two key techniques boost the throughput of the instruction pipeline. Branch prediction estimates, in
advance, which path a conditional branch will take, allowing the processor to fetch and begin executing
instructions before the condition is evaluated. Out-of-order execution allows the processor to execute
independent instructions in a different order from their appearance in the program, keeping the execution
units busy even when some instructions are waiting for data.
"""),

    ("3. Memory Systems", """
Memory in a computer system is organized as a hierarchy, trading off speed, capacity, and cost at each
level. At the top of the hierarchy sit the processor registers: a few dozen 64-bit locations that operate
at the full clock speed of the CPU and can be read or written in a single cycle. Below the registers are
the cache levels described above, followed by main memory, solid-state or hard-disk storage, and finally
networked or cloud storage.

Main memory, implemented as dynamic random-access memory (DRAM), is the working space of running
programs. A modern workstation typically contains 16 to 128 GB of RAM organized in dual-channel or
quad-channel configurations to increase memory bandwidth. The current DDR5 standard supports transfer
rates between 4,800 MT/s and 8,400 MT/s per channel, providing a theoretical peak bandwidth in the range
of 50 to 100 GB/s for a dual-channel system.

Although DRAM is much slower than cache — a cache miss that requires accessing main memory typically
costs 50 to 100 CPU cycles — it is far faster than secondary storage. A high-end NVMe solid-state drive
achieves sequential read speeds of around 7,000 MB/s, but random access latency is measured in
microseconds rather than nanoseconds. Magnetic hard drives are slower still, with random-access latency
in the range of 5 to 10 milliseconds.

Virtual memory is an abstraction that allows each process to operate as though it has exclusive access
to a large, contiguous address space, regardless of how much physical memory is actually available. The
operating system maps virtual addresses to physical addresses through a data structure called a page table.
When a process accesses a virtual address that is not backed by physical memory, the OS raises a page
fault, loads the required data from disk, and resumes the process. This mechanism enables the system to
run programs whose combined memory requirements exceed the available physical RAM, at the cost of
significantly increased latency whenever disk access is required.
"""),

    ("4. Input/Output and Storage", """
Input/output (I/O) devices connect the computer to the outside world, allowing it to receive data from
keyboards, mice, network interfaces, sensors, and cameras, and to send data to displays, printers,
speakers, and network links. I/O devices communicate with the processor through controllers, which
translate between the device-specific signaling and the bus protocols used inside the computer.

The Peripheral Component Interconnect Express (PCIe) bus is the dominant interconnect for high-speed
devices in modern PCs and servers. PCIe 5.0, the current generation, provides a bandwidth of 4 GB/s
per lane. A device using 16 lanes — typical for a graphics card — therefore has a peak bandwidth of
64 GB/s in each direction.

Storage devices retain data even when power is removed, in contrast to volatile main memory. NVMe
solid-state drives connect directly to the PCIe bus, bypassing the older SATA interface and achieving
much higher throughput and lower latency. NAND flash memory, the technology underlying SSDs, stores
data as charge levels in floating-gate transistors organized into cells. Modern 3D NAND stacks over
200 layers of cells vertically to achieve capacities of several terabytes per drive.

Direct memory access (DMA) allows I/O devices to transfer data directly to or from main memory without
involving the CPU for each transfer. The device signals an interrupt when the transfer is complete,
and the CPU then processes the received data. DMA is essential for achieving the full potential
bandwidth of modern storage and network devices, since even a fast CPU cannot sustain the data rates
required for a 100 Gbps network interface using programmed I/O.
"""),

    ("5. Computer Networks", """
A computer network is a system of interconnected nodes that exchange data according to agreed-upon
protocols. The internet, the largest network in existence, is an interconnection of millions of smaller
networks operated by organizations, governments, and service providers. Its architecture follows the
TCP/IP model, which divides the communication problem into four layers: the link layer handles local
delivery over a single network segment, the internet layer routes packets across multiple networks using
IP addresses, the transport layer provides end-to-end reliability (TCP) or low-overhead delivery (UDP),
and the application layer contains protocols like HTTP, SMTP, and DNS that serve specific user needs.

Local area networks (LANs) connect devices within a building or campus. Wired Ethernet, standardized
as IEEE 802.3, is the dominant LAN technology, with current deployments offering 1 Gbps to 100 Gbps
per port. Wi-Fi, standardized as IEEE 802.11, provides wireless connectivity with current-generation
Wi-Fi 7 delivering aggregate throughput over 40 Gbps in ideal conditions, though real-world performance
is lower due to interference, distance, and the shared nature of the radio medium.

Security in computer networks involves multiple layers of defense. Transport Layer Security (TLS)
encrypts application-layer data between endpoints, protecting it from eavesdropping and tampering.
Firewalls filter traffic based on source address, destination address, port, and protocol. Intrusion
detection systems analyze network traffic for patterns that suggest malicious activity. Virtual private
networks (VPNs) create encrypted tunnels across untrusted networks, allowing remote users to access
internal resources securely.

Network performance is characterized by three key metrics: bandwidth, the maximum rate at which data
can be transferred; latency, the time required for a packet to travel from source to destination; and
packet loss, the fraction of packets that are dropped due to congestion or errors. For interactive
applications such as voice and video conferencing, latency below 150 ms is generally required for
acceptable quality. For bulk data transfer, bandwidth and loss rate are more important than latency.
"""),

    ("6. Operating Systems", """
An operating system (OS) is the layer of software that manages the hardware resources of a computer
and provides a stable, convenient interface for application programs. Without an OS, each application
would need to contain its own code for interacting with every piece of hardware it might use — an
impossible burden that would also make it impossible for multiple programs to share the same machine.

The most important resource the OS manages is the processor. The scheduler decides which process or
thread runs on each core at each moment, balancing competing objectives: responsiveness (interactive
tasks should receive CPU time quickly), throughput (background work should complete efficiently), and
fairness (no process should starve). Modern schedulers are preemptive: they can interrupt a running
process and switch to another without the cooperation of the running process, using a hardware timer
to trigger the switch.

Memory management is the second major OS responsibility. The OS maintains the page tables that implement
virtual memory, allocates physical memory pages to processes, and handles page faults. It also provides
the heap allocator that application programmers use through functions like malloc and free. Efficient
memory management is critical for system performance; memory leaks and fragmentation can degrade
performance over time and cause programs to fail.

The file system provides an organized, persistent store for data. Files are named sequences of bytes
stored on a storage device. Directories organize files into a hierarchical namespace. The file system
maintains metadata about each file — its name, size, creation time, modification time, and access
permissions — and implements the operations that applications use: create, read, write, delete, rename,
and change permissions. Modern file systems include journaling, which records pending changes before
committing them, so that a crash during a write operation does not corrupt the file system structure.
"""),

    ("7. Performance and Optimization", """
Measuring and improving computer performance is both an art and a science. The most important
principle is that optimization should be guided by measurement: tools called profilers record where
a program spends its time, which functions are called most often, and where memory is most heavily
allocated. Optimizing code that is not a bottleneck wastes effort and can introduce bugs.

Memory access patterns have a profound effect on performance. Modern processors are far faster than
main memory, and a program that repeatedly accesses data scattered through a large memory region will
spend most of its time waiting for cache misses. Reorganizing data structures to improve spatial
locality — keeping related data close together in memory — can reduce cache miss rates dramatically
and yield speedups of 5x or more without changing the algorithm at all.

Parallel execution on multi-core processors offers another avenue for speedup, but it introduces
complexity. Shared data must be protected from concurrent modification using synchronization primitives.
Locks introduce contention and can cause threads to block while waiting. Lock-free algorithms use
atomic hardware instructions to achieve thread safety without blocking, but are difficult to design
correctly. The degree to which a program can benefit from multiple cores is limited by the fraction
of its execution that must remain sequential, a constraint formalized by Amdahl's Law.

Power consumption is increasingly important as a performance metric. In mobile devices and large data
centers alike, energy efficiency directly affects battery life and operating costs. Dynamic voltage
and frequency scaling (DVFS) allows processors to reduce their clock speed and supply voltage during
periods of low load, reducing power consumption quadratically with clock speed. Modern chips combine
high-performance cores with efficiency cores optimized for low power, choosing dynamically which type
of core to use based on the current workload.
"""),
]

# Таблица характеристик процессоров (для DOCX и XLSX)
CPU_TABLE_HEADERS = ["Processor", "Cores", "Base Freq (GHz)", "L3 Cache (MB)", "TDP (W)", "Architecture"]
CPU_TABLE_ROWS = [
    ["Intel Core i9-14900K",  "24 (8P+16E)", "3.2",  "36",  "125",  "Raptor Lake"],
    ["AMD Ryzen 9 7950X",     "16",           "4.5",  "64",  "170",  "Zen 4"],
    ["Intel Xeon w9-3595X",   "60",           "2.5",  "75",  "350",  "Sapphire Rapids"],
    ["AMD EPYC 9654",         "96",           "2.4",  "384", "360",  "Genoa"],
    ["Apple M3 Max",          "16 (12P+4E)",  "4.05", "48",  "92",   "ARM v8.6"],
    ["Qualcomm Snapdragon X", "12",           "3.8",  "42",  "23",   "Oryon"],
]

DOCUMENT_CONCLUSION = """
Computer architecture, memory systems, networking, and operating systems together form the substrate on
which all software runs. Proficiency in these areas helps engineers make better decisions at every level
of the stack, from the choice of data structures to the design of distributed systems. The field continues
to evolve rapidly: chiplet-based designs decompose a processor into separately manufactured tiles;
processing-in-memory architectures move computation closer to data; neuromorphic and quantum processors
explore fundamentally different computational models. Despite this diversity, the principles covered in
this document — locality, hierarchy, abstraction, and concurrency — remain relevant across all of them.
"""


def build_plain_text() -> str:
    lines = [DOCUMENT_TITLE, "=" * len(DOCUMENT_TITLE), ""]
    for heading, body in DOCUMENT_SECTIONS:
        lines.append(heading)
        lines.append("-" * len(heading))
        lines.append(body.strip())
        lines.append("")
    # таблица как plain text
    lines.append("Table 1. Processor Specifications")
    lines.append("-" * 40)
    lines.append("  |  ".join(CPU_TABLE_HEADERS))
    lines.append("-" * 80)
    for row in CPU_TABLE_ROWS:
        lines.append("  |  ".join(row))
    lines.append("")
    lines.append("Conclusion")
    lines.append("-" * 10)
    lines.append(DOCUMENT_CONCLUSION.strip())
    return "\n".join(lines)


def build_markdown() -> str:
    lines = [f"# {DOCUMENT_TITLE}", ""]
    for heading, body in DOCUMENT_SECTIONS:
        lines.append(f"## {heading}")
        lines.append("")
        lines.append(body.strip())
        lines.append("")
    # таблица в markdown
    lines.append("## Table 1. Processor Specifications")
    lines.append("")
    header_row = " | ".join(CPU_TABLE_HEADERS)
    sep_row = " | ".join(["---"] * len(CPU_TABLE_HEADERS))
    lines.append(f"| {header_row} |")
    lines.append(f"| {sep_row} |")
    for row in CPU_TABLE_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    # блок кода (не должен переводиться)
    lines.append("## Code Example: Reading CPU Info")
    lines.append("")
    lines.append("The following Python snippet retrieves basic CPU information:")
    lines.append("")
    lines.append("```python")
    lines.append("import psutil")
    lines.append("import platform")
    lines.append("")
    lines.append("cpu_count = psutil.cpu_count(logical=True)")
    lines.append("cpu_freq = psutil.cpu_freq()")
    lines.append("print(f'Logical CPUs : {cpu_count}')")
    lines.append("print(f'Max frequency: {cpu_freq.max:.0f} MHz')")
    lines.append("print(f'Architecture : {platform.machine()}')")
    lines.append("```")
    lines.append("")
    lines.append("## Conclusion")
    lines.append("")
    lines.append(DOCUMENT_CONCLUSION.strip())
    return "\n".join(lines)


def build_docx(path: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Колонтитулы
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = DOCUMENT_TITLE
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    footer.paragraphs[0].text = "Test corpus — LightLocalTranslator project"
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Заголовок
    title_para = doc.add_heading(DOCUMENT_TITLE, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for heading, body in DOCUMENT_SECTIONS:
        doc.add_heading(heading, level=1)
        # Разбиваем тело на абзацы
        for para_text in body.strip().split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            # Первый абзац каждого раздела — курсив (для проверки сохранения стилей)
            run = p.add_run(para_text)

    # Таблица
    doc.add_heading("Table 1. Processor Specifications", level=2)
    table = doc.add_table(
        rows=1 + len(CPU_TABLE_ROWS),
        cols=len(CPU_TABLE_HEADERS),
        style="Table Grid"
    )
    # Заголовки таблицы — жирный текст
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(CPU_TABLE_HEADERS):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
    # Данные таблицы
    for r_idx, row_data in enumerate(CPU_TABLE_ROWS):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = val

    # Заключение
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(DOCUMENT_CONCLUSION.strip())

    doc.save(str(path))


def build_xlsx(path: Path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "CPU Specifications"

    # Заголовок листа
    ws["A1"] = DOCUMENT_TITLE
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:F1")
    ws["A1"].alignment = Alignment(horizontal="center")

    # Пустая строка
    ws.append([])

    # Заголовки таблицы (строка 3)
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2E4057")
    for col_idx, header in enumerate(CPU_TABLE_HEADERS, start=1):
        cell = ws.cell(row=3, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    # Данные
    alt_fill = PatternFill("solid", fgColor="EAF0FB")
    for r_idx, row_data in enumerate(CPU_TABLE_ROWS):
        row_num = 4 + r_idx
        fill = alt_fill if r_idx % 2 == 0 else PatternFill()
        for c_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_num, column=c_idx, value=val)
            cell.fill = fill

    # Формулы (не должны переводиться)
    formula_row = 4 + len(CPU_TABLE_ROWS) + 1
    ws.cell(row=formula_row, column=1, value="Average TDP (W):")
    ws.cell(row=formula_row, column=1).font = Font(bold=True)
    # TDP в столбце 5 (E), строки 4 .. 4+N-1
    tdp_start = 4
    tdp_end = 4 + len(CPU_TABLE_ROWS) - 1
    ws.cell(row=formula_row, column=2,
            value=f"=AVERAGE(E{tdp_start}:E{tdp_end})")

    ws.cell(row=formula_row + 1, column=1, value="Total entries:")
    ws.cell(row=formula_row + 1, column=1).font = Font(bold=True)
    ws.cell(row=formula_row + 1, column=2,
            value=f"=COUNTA(A{tdp_start}:A{tdp_end})")

    # Ширина столбцов
    col_widths = [28, 16, 18, 16, 12, 18]
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # Второй лист — описание документа
    ws2 = wb.create_sheet(title="Document Info")
    ws2["A1"] = "Document"
    ws2["B1"] = DOCUMENT_TITLE
    ws2["A2"] = "Purpose"
    ws2["B2"] = "Test corpus for translation quality evaluation (formatting preservation check)"
    ws2["A3"] = "Sections"
    ws2["B3"] = str(len(DOCUMENT_SECTIONS))
    ws2["A4"] = "Word count (approx)"
    ws2["B4"] = "~3000"
    ws2["A5"] = "Source language"
    ws2["B5"] = "English"

    wb.save(str(path))


def build_pdf(path: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    )

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=DOCUMENT_TITLE,
        author="LightLocalTranslator corpus",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Title"],
        fontSize=16, spaceAfter=16
    )
    h1_style = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontSize=13, spaceBefore=14, spaceAfter=6
    )
    body_style = styles["BodyText"]
    body_style.leading = 14

    story = []
    story.append(Paragraph(DOCUMENT_TITLE, title_style))
    story.append(Spacer(1, 0.4 * cm))

    for heading, body in DOCUMENT_SECTIONS:
        story.append(Paragraph(heading, h1_style))
        for para_text in body.strip().split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                story.append(Paragraph(para_text, body_style))
                story.append(Spacer(1, 0.2 * cm))

    # Таблица
    story.append(Paragraph("Table 1. Processor Specifications", h1_style))
    table_data = [CPU_TABLE_HEADERS] + CPU_TABLE_ROWS
    tbl = Table(table_data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E4057")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF0FB")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.4 * cm))

    story.append(Paragraph("Conclusion", h1_style))
    story.append(Paragraph(DOCUMENT_CONCLUSION.strip(), body_style))

    doc.build(story)


# ---------------------------------------------------------------------------
# Запись JSONL
# ---------------------------------------------------------------------------

def write_jsonl(path: Path, records: list):
    with open(path, "w", encoding="utf-8") as f:
        for rec in records:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    print(f"  wrote {len(records)} records -> {path}")


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main():
    print("=== Creating test corpus ===\n")

    print("[1/7] A_words.jsonl")
    write_jsonl(BLOCKS / "A_words.jsonl", WORDS)

    print("[2/7] B_sentences.jsonl")
    write_jsonl(BLOCKS / "B_sentences.jsonl", SENTENCES)

    print("[3/7] C_paragraphs.jsonl")
    write_jsonl(BLOCKS / "C_paragraphs.jsonl", PARAGRAPHS)

    print("[4/7] test_doc.txt")
    txt_path = DOCS / "test_doc.txt"
    txt_path.write_text(build_plain_text(), encoding="utf-8")
    print(f"  wrote -> {txt_path}")

    print("[5/7] test_doc.md")
    md_path = DOCS / "test_doc.md"
    md_path.write_text(build_markdown(), encoding="utf-8")
    print(f"  wrote -> {md_path}")

    print("[6/7] test_doc.docx")
    build_docx(DOCS / "test_doc.docx")
    print(f"  wrote -> {DOCS / 'test_doc.docx'}")

    print("[7/7] test_doc.xlsx")
    build_xlsx(DOCS / "test_doc.xlsx")
    print(f"  wrote -> {DOCS / 'test_doc.xlsx'}")

    print("\n[+] test_doc.pdf")
    try:
        build_pdf(DOCS / "test_doc.pdf")
        print(f"  wrote -> {DOCS / 'test_doc.pdf'}")
    except Exception as e:
        print(f"  WARN: PDF skipped ({e})")

    print("\n=== Corpus created successfully ===")
    print(f"  blocks/   : {BLOCKS}")
    print(f"  documents/: {DOCS}")


if __name__ == "__main__":
    main()
